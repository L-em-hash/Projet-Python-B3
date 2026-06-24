# PROJET PYTHON AVANCÉ - PARTIE 2
# Rapport Word automatique depuis Project Gutenberg
# Dépendances : pip install requests matplotlib pillow python-docx

import requests
import re
import os
import io
from collections import Counter

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

LIVRE_URL = "https://www.gutenberg.org/cache/epub/135/pg135.txt"
IMAGE_URL = "https://upload.wikimedia.org/wikipedia/commons/thumb/e/e6/Les_Miserables_poster.jpg/220px-Les_Miserables_poster.jpg"

LOGO_PATH = "data/logo.png"
OUTPUT_WORD = "data/rapport_livre.docx"
GRAPH_PATH = "data/graphique_paragraphes.png"
IMAGE_PATH = "data/image_livre.png"

MON_NOM = "LAWSON-LARTEGO Nadou Emmanuella"


def telecharger_livre(url):
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    try:
        return r.content.decode("utf-8")
    except UnicodeDecodeError:
        return r.content.decode("latin-1")


def extraire_metadonnees(texte):
    titre = "Titre inconnu"
    auteur = "Auteur inconnu"

    for ligne in texte.splitlines()[:50]:
        if ligne.startswith("Title:"):
            titre = ligne.replace("Title:", "").strip()
        if ligne.startswith("Author:"):
            auteur = ligne.replace("Author:", "").strip()

    return titre, auteur


def extraire_premier_chapitre(texte):
    patron = re.compile(
        r'\n(CHAPTER\s+[IVXLC\d]+|CHAPITRE\s+[IVXLC\d]+|LIVRE\s+[IVXLC\d]+)',
        re.IGNORECASE
    )

    matches = list(patron.finditer(texte))

    if len(matches) >= 2:
        return texte[matches[0].start():matches[1].start()].strip()
    elif len(matches) == 1:
        return texte[matches[0].start():matches[0].start()+5000].strip()

    idx = texte.find("*** START OF")
    return texte[idx+200:idx+5200] if idx != -1 else texte[1000:6000]


def compter_mots_paragraphes(chapitre):
    paragraphes = [
        p.strip() for p in re.split(r'\n\s*\n', chapitre)
        if p.strip()
    ]

    nb_mots = []

    for p in paragraphes:
        mots = len(p.split())
        if mots > 0:
            nb_mots.append((mots // 10) * 10)

    return nb_mots, paragraphes


def statistiques(nb_mots, paragraphes):
    total_mots = sum(len(p.split()) for p in paragraphes)

    return {
        "nb_paragraphes": len(nb_mots),
        "total_mots": total_mots,
        "moy_mots": round(total_mots / len(paragraphes), 1) if paragraphes else 0,
        "min_mots": min(len(p.split()) for p in paragraphes) if paragraphes else 0,
        "max_mots": max(len(p.split()) for p in paragraphes) if paragraphes else 0,
    }


def creer_graphique(nb_mots, titre_livre, chemin):
    compteur = Counter(nb_mots)

    x = sorted(compteur.keys())
    y = [compteur[v] for v in x]

    plt.figure(figsize=(10, 5))
    plt.bar(x, y)
    plt.title(f"Distribution des paragraphes - {titre_livre}")
    plt.xlabel("Mots arrondis à la dizaine")
    plt.ylabel("Nombre de paragraphes")
    plt.tight_layout()
    plt.savefig(chemin)
    plt.close()


def telecharger_image(url, chemin):
    try:
        r = requests.get(url, timeout=20)
        r.raise_for_status()

        img = Image.open(io.BytesIO(r.content)).convert("RGB")

        w, h = img.size
        marge = min(w, h) // 6

        img = img.crop((marge, marge, w-marge, h-marge))
        img = img.resize((400, 500))

        img.save(chemin)

    except Exception:
        img = Image.new("RGB", (400, 500))
        img.save(chemin)


def creer_logo(chemin):
    if not os.path.exists(chemin):
        from PIL import ImageDraw

        img = Image.new("L", (200, 200), color=255)
        draw = ImageDraw.Draw(img)

        draw.ellipse([20, 20, 180, 180], outline=0, width=8)
        draw.text((60, 85), "LOGO", fill=0)

        img.save(chemin)


def coller_logo(image_chemin, logo_chemin):
    img = Image.open(image_chemin).convert("RGBA")
    logo = Image.open(logo_chemin).convert("L")

    logo = ImageOps.invert(logo)
    logo = logo.convert("RGBA")
    logo = logo.rotate(25, expand=True)
    logo = logo.resize((100, 100))

    img.paste(logo, (10, 10), logo)
    img.convert("RGB").save(image_chemin)


def generer_word(
    titre,
    auteur,
    stats,
    image_chemin,
    graph_chemin,
    chapitre,
    mon_nom,
    chemin_sortie
):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    p = doc.add_paragraph(titre)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists(image_chemin):
        doc.add_picture(image_chemin, width=Inches(3))

    doc.add_paragraph(f"Auteur : {auteur}")
    doc.add_paragraph(f"Rapport rédigé par : {mon_nom}")

    doc.add_page_break()

    doc.add_heading("Analyse du premier chapitre", level=1)

    if os.path.exists(graph_chemin):
        doc.add_picture(graph_chemin, width=Inches(5))

    doc.add_heading("Description et statistiques", level=2)

    doc.add_paragraph(
        f'Ce rapport analyse le premier chapitre de "{titre}", '
        f'œuvre de {auteur}.'
    )

    resume = (
        f"Le premier chapitre contient {stats['nb_paragraphes']} paragraphes "
        f"pour un total de {stats['total_mots']} mots. "
        f"La longueur moyenne est de {stats['moy_mots']} mots par paragraphe. "
        f"Le paragraphe le plus court contient {stats['min_mots']} mots "
        f"et le plus long {stats['max_mots']} mots."
    )

    doc.add_paragraph(resume)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    table.rows[0].cells[0].text = "Indicateur"
    table.rows[0].cells[1].text = "Valeur"

    donnees = [
        ("Nombre de paragraphes", stats["nb_paragraphes"]),
        ("Nombre total de mots", stats["total_mots"]),
        ("Moyenne", stats["moy_mots"]),
        ("Minimum", stats["min_mots"]),
        ("Maximum", stats["max_mots"]),
    ]

    for k, v in donnees:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    doc.add_heading("Extrait", level=2)
    doc.add_paragraph(chapitre[:800])

    doc.save(chemin_sortie)


def main():
    texte = telecharger_livre(LIVRE_URL)

    titre, auteur = extraire_metadonnees(texte)
    chapitre = extraire_premier_chapitre(texte)

    nb_mots, paragraphes = compter_mots_paragraphes(chapitre)
    stats = statistiques(nb_mots, paragraphes)

    creer_graphique(nb_mots, titre, GRAPH_PATH)

    telecharger_image(IMAGE_URL, IMAGE_PATH)
    creer_logo(LOGO_PATH)
    coller_logo(IMAGE_PATH, LOGO_PATH)

    generer_word(
        titre,
        auteur,
        stats,
        IMAGE_PATH,
        GRAPH_PATH,
        chapitre,
        MON_NOM,
        OUTPUT_WORD
    )

    print("Rapport généré :", OUTPUT_WORD)


if __name__ == "__main__":
    main()
